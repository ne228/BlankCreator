from docx import Document
from datetime import datetime
import os
from docx.shared import Cm
from docx.shared import Pt
from docxcompose.composer import Composer
import random
# from docx2pdf import convert

# var inits

len_sub_str = 35
directory_save = "docs"

def split_text_by_length(text):
    words = text.split()  # Разбиваем строку на слова
    result = []           # Результирующий список подстрок
    current_line = []     # Текущая подстрока, которая собирается

    for word in words:
        # Проверяем, помещается ли слово в текущую строку
        if sum(len(w) for w in current_line) + len(word) + len(current_line) <= len_sub_str:
            current_line.append(word)
        else:
            # Если не помещается, добавляем текущую строку в результат и начинаем новую строку
            result.append(' '.join(current_line))
            current_line = [word]

    # Добавляем последнюю строку, если она не пуста
    if current_line:
        result.append(' '.join(current_line))

    return result

def get_datetime_filename(prefix = "doc"):
    current_time = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    file_name = f"{prefix}_{current_time}.docx"
    return file_name
    
def replace_keywords_in_paragraph(paragraph, replacements):
    """Заменяет ключевые слова в абзаце с сохранением формата."""
    for key, value in replacements.items():
        for run in paragraph.runs:
            if key in run.text:
                # Обработка place
                
                if key == "place":
                    # Делим на подстроки
                    sub_strs = split_text_by_length(replacements["place"]) 
                      
                    # если подстрок больше 1, то оставшуюсую часть place вставляем в town
                    if len(sub_strs) > 1:
                        first, *second = sub_strs
                        second_str = ''.join(map(str, second))
                        town_str = replacements["town"]
                        replacements["town"] = f"{second_str}, {town_str}"
                        run.text = run.text.replace(key, first)
                        continue
                
                run.text = run.text.replace(key, value)

def replace_keywords_in_table(table, replacements):
    """Заменяет ключевые слова во всех ячейках таблицы."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                replace_keywords_in_paragraph(paragraph, replacements)

def create_individual_document(template_path, replacements, output_path, index):
    """Создает индивидуальный документ на основе шаблона и заменяет ключевые слова."""
    doc = Document(template_path)
    for table in doc.tables:
        replace_keywords_in_table(table, replacements)        

    # Каждому второму документу вставляем в конце разрыва раздела
    if (index + 1) % 2 == 0:
        # Добавляем разрыв раздела в конец документа
        doc.add_section()
    else:
        # Добавляем параграф для линии
        paragraph = doc.add_paragraph()
        run = paragraph.add_run('_' * 150)  # создаем линию из символов подчеркивания
        run.font.size = Pt(10)  # уменьшаем высоту линии
        run.font.color.rgb = None  # устанавливаем черный цвет линии

        
    doc.save(output_path)

def create_documents_for_data(template_path, data_list):
    """Создает документ для каждого словаря в списке данных и возвращает массив путей."""
    document_paths = []


    for i, data in enumerate(data_list):
        # Генерация уникального имени файла с текущей датой и временем
        file_name = get_datetime_filename(f"{directory_save}/doc_{i}")
        # print(file_name)
        create_individual_document(template_path, data, file_name, i)
        document_paths.append(file_name)
    
    return document_paths



def set_document_margins(path):
    # Открываем документ
    doc = Document(path)
    
    # Устанавливаем поля на 0.3 см
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.3)
        section.bottom_margin = Cm(0.3)
        section.left_margin = Cm(0.6)
        section.right_margin = Cm(0.3)
    
    # Сохраняем изменения в новом документе
    
    output_path = path  # Создаем новое имя файла
    doc.save(output_path)
    return output_path



def merge_documents(output_path, *input_paths):
    merged_doc = Document()
    composer = Composer(merged_doc)
    for path in input_paths:
        print(f"Compose item {path}")
        composer.append(Document(path))
    composer.save(output_path)

def convert_docx_to_pdf(docx_file):
    # Проверяем, что файл существует
    if not os.path.exists(docx_file):
        raise FileNotFoundError(f"Файл {docx_file} не найден.")

    # Конвертируем файл .docx в .pdf с помощью LibreOffice
    subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', docx_file], check=True)

    # Формируем путь к выходному PDF файлу
    base = os.path.splitext(docx_file)[0]  # Убираем расширение .docx
    pdf_file = f"{base}.pdf"

    # Проверяем, создан ли файл .pdf
    if not os.path.exists(pdf_file):
        raise FileNotFoundError(f"Не удалось создать PDF файл {pdf_file}.")

    return pdf_file
                
# Пример данных
# data_list = [
#     {
#         "rank": "Старший лейтенант",
#         "name": "Иван Иванов",
#         "date_birth": "01.01.1990",
#         "duty": "Офицер",
#         "date_pr": "01.01.2024",
#         "num_pr": "123456",
#         "date_enrollment": "15.03.2015",
#         "trm": "5",
#         "date_end": "30.01.2024",
#         "place": "Очень длинная длинная длинная длинная длинная военная часть №123",
#         "town": "Москва"
#     },
#     {
#         "rank": "Лейтенант",
#         "name": "Петр Петров",
#         "date_birth": "05.05.1991",
#         "duty": "Инженер",
#         "date_pr": "15.02.2024",
#         "num_pr": "654321",
#         "date_enrollment": "10.04.2016",
#         "trm": "3",
#         "date_end": "15.02.2025",
#         "place": "Военная часть №456",
#         "town": "Санкт-Петербург"
#     }
# ]


def create_blank_docx(data):

    if not os.path.exists(directory_save):
        os.makedirs(directory_save)

    id = random.randint(0, 100)
    print(f"Start create blank: {id}")
    path_template = "template.docx"
    path_merged_doc = get_datetime_filename(f"{directory_save}/res_{id}")

    # create documetns
    document_files = create_documents_for_data(path_template, data)
    print(f"Documents saved: {document_files}")

    # merge documents
    merge_documents(path_merged_doc, *document_files)
    print(f"Merge success: {path_merged_doc}")
    for doc in document_files:
        os.remove(doc)

    # set margin page
    result_path = set_document_margins(path_merged_doc)
    print(f"Blank saved as: {result_path}")
    return result_path

def create_blank_pdf(data):
    path_docx = create_blank_docx(data)
    file_name = os.path.splitext(path_docx)[0]

    # Конвертируем файл .docx в .pdf
    convert(path_docx)
    output_file = file_name + ".pdf"
    return
 